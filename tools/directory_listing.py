#!/usr/bin/env python3
"""
Generate a directory listing (tree + files) and save it to a Word document.

Usage:
    python directory_listing.py /path/to/folder
    python directory_listing.py /path/to/folder -o listing.docx

Requires: pip install python-docx
"""

import argparse
import sys
from pathlib import Path
from tkinter import Tk, filedialog, messagebox

from docx import Document
from docx.shared import Pt


def walk_tree(path: Path, prefix: str = "", entries: list | None = None) -> list:
    """Recursively build a list of (line, is_dir) tuples that form the tree."""
    if entries is None:
        entries = []

    try:
        # Folders first, then files, both alphabetically (case-insensitive)
        items = sorted(
            path.iterdir(),
            key=lambda p: (not p.is_dir(), p.name.lower()),
        )
    except PermissionError:
        entries.append((prefix + "└── [permission denied]", False))
        return entries

    for i, item in enumerate(items):
        is_last = i == len(items) - 1
        connector = "└── " if is_last else "├── "
        name = item.name + ("/" if item.is_dir() else "")
        entries.append((prefix + connector + name, item.is_dir()))

        if item.is_dir():
            extension = "    " if is_last else "│   "
            walk_tree(item, prefix + extension, entries)

    return entries


def create_document(root: Path, entries: list, output_path: Path) -> None:
    """Write the entries into a .docx file."""
    doc = Document()

    # Title and summary
    doc.add_heading(f"Directory Listing: {root}", level=1)

    n_dirs = sum(1 for _, is_dir in entries if is_dir)
    n_files = sum(1 for _, is_dir in entries if not is_dir)
    summary = doc.add_paragraph()
    summary.add_run(
        f"{n_dirs} folder{'s' if n_dirs != 1 else ''}, "
        f"{n_files} file{'s' if n_files != 1 else ''}"
    ).italic = True

    # Root of the tree
    root_para = doc.add_paragraph()
    root_run = root_para.add_run(f"{root.name or str(root)}/")
    root_run.font.name = "Consolas"
    root_run.font.size = Pt(10)
    root_run.bold = True

    # Tree body — monospaced so connectors line up cleanly
    for line, is_dir in entries:
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        if is_dir:
            run.bold = True

    doc.save(output_path)


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Generate a directory listing and save it as a Word document."
    )
    parser.add_argument("folder", help="Folder to list")
    parser.add_argument(
        "-o",
        "--output",
        default="directory_listing.docx",
        help="Output .docx file (default: directory_listing.docx)",
    )
    args = parser.parse_args()

    root = Path(args.folder).resolve()
    if not root.is_dir():
        print(f"Error: {root} is not a directory", file=sys.stderr)
        return 1

    entries = walk_tree(root)
    output_path = Path(args.output).resolve()
    create_document(root, entries, output_path)

    print(f"Wrote {len(entries)} entries to {output_path}")
    return 0


def run_directory_listing_gui() -> str:
    root = Tk()
    root.withdraw()

    folder = filedialog.askdirectory(title="选择要导出的目录")
    if not folder:
        return "已取消：目录树导出。"

    output_path = filedialog.asksaveasfilename(
        title="保存目录清单",
        defaultextension=".docx",
        initialfile="directory_listing.docx",
        filetypes=[("Word 文档", "*.docx")],
    )
    if not output_path:
        return "已取消：目录树导出。"

    try:
        root_path = Path(folder).resolve()
        entries = walk_tree(root_path)
        create_document(root_path, entries, Path(output_path).resolve())
        return f"已导出 {len(entries)} 条目录记录：{output_path}"
    except Exception as exc:
        messagebox.showerror("导出失败", str(exc))
        return f"目录树导出失败：{exc}"


if __name__ == "__main__":
    sys.exit(main())
