import os
from docx import Document

def replace_text_in_paragraph(paragraph, old_text, new_text):
    """替换段落中的文字，包括 Run 对象"""
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)

def replace_text_in_table(table, old_text, new_text):
    """替换表格中的文字"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, old_text, new_text)

def replace_text_in_headers_footers(doc, old_text, new_text):
    """替换页眉和页脚中的文字"""
    for section in doc.sections:
        # 处理页眉
        header = section.header
        for paragraph in header.paragraphs:
            replace_text_in_paragraph(paragraph, old_text, new_text)
        # 处理页脚
        footer = section.footer
        for paragraph in footer.paragraphs:
            replace_text_in_paragraph(paragraph, old_text, new_text)

def replace_text_in_docx(docx_file, old_text, new_text):
    """替换文档中的文字，包括正文、表格、页眉和页脚"""
    doc = Document(docx_file)
    # 替换正文段落内容
    for para in doc.paragraphs:
        replace_text_in_paragraph(para, old_text, new_text)
    # 替换表格中的内容
    for table in doc.tables:
        replace_text_in_table(table, old_text, new_text)
    # 替换页眉和页脚中的内容
    replace_text_in_headers_footers(doc, old_text, new_text)
    # 保存修改后的文档
    doc.save(docx_file)

def search_and_replace(root_folder, old_text, new_text):
    """递归处理文件夹中的所有 .docx 文件"""
    for root, dirs, files in os.walk(root_folder):
        for file in files:
            if file.endswith('.docx'):
                docx_file = os.path.join(root, file)
                print(f"Processing {docx_file}")
                try:
                    replace_text_in_docx(docx_file, old_text, new_text)
                    print(f"Replaced text in {docx_file}")
                except Exception as e:
                    print(f"Failed to process {docx_file}: {e}")

# 使用示例
root_folder = r'D:\OneDrive\3. Business\3.1 产品文件\2.4 产品文件\0. 标准文件\3.1 蛇毒肽 粉末 EN'  # 替换为你的文件夹路径
old_text = '1st Floor, B5 Building, University Park, No. 16'  # 要替换的旧文本
new_text = '1st Floor, Building B5,  No.16 Haichuan Road,'  # 新文本

search_and_replace(root_folder, old_text, new_text)
