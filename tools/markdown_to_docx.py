import os
from tkinter import Tk, filedialog, messagebox


def _load_conversion_dependencies():
    try:
        import markdown
        from bs4 import BeautifulSoup
        from docx import Document
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "缺少 Markdown 转 Word 依赖，请先安装：pip install markdown beautifulsoup4 python-docx"
        ) from exc
    return markdown, BeautifulSoup, Document

def select_markdown_file():
    root = Tk()
    root.withdraw()
    file_path = filedialog.askopenfilename(
        title="请选择一个 Markdown 文件",
        filetypes=[("Markdown 文件", "*.md *.markdown"), ("所有文件", "*.*")]
    )
    return file_path

def markdown_to_word(input_file, output_file):
    try:
        markdown, BeautifulSoup, Document = _load_conversion_dependencies()

        # Step 1: 读取 Markdown 文件
        with open(input_file, 'r', encoding='utf-8') as f:
            markdown_text = f.read()
        
        # Step 2: Markdown -> HTML（加入扩展）
        html_body = markdown.markdown(markdown_text, extensions=['extra'])
        html = f"<html><body>{html_body}</body></html>"

        # Step 3: 使用 BeautifulSoup 解析 HTML
        soup = BeautifulSoup(html, 'html.parser')
        body = soup.body

        # Step 4: 创建 Word 文档
        doc = Document()
        doc.add_heading('Markdown 转换结果', level=0)

        # Step 5: 遍历 HTML 元素，写入 Word
        for elem in body.children:
            if elem.name is None:
                continue  # 忽略空行
            elif elem.name == 'h1':
                doc.add_heading(elem.get_text(), level=1)
            elif elem.name == 'h2':
                doc.add_heading(elem.get_text(), level=2)
            elif elem.name == 'h3':
                doc.add_heading(elem.get_text(), level=3)
            elif elem.name == 'p':
                doc.add_paragraph(elem.get_text())
            elif elem.name == 'ul':
                for li in elem.find_all('li'):
                    doc.add_paragraph("• " + li.get_text(), style='ListBullet')
            elif elem.name == 'ol':
                for li in elem.find_all('li'):
                    doc.add_paragraph(li.get_text(), style='ListNumber')
            else:
                doc.add_paragraph(elem.get_text())

        # Step 6: 保存 Word 文档
        doc.save(output_file)
        print(f"\n✅ 成功生成 Word 文件：{output_file}")
    except Exception as e:
        try:
            messagebox.showerror("转换失败", str(e))
        except Exception:
            pass
        print(f"❌ 转换失败: {str(e)}")

def run_markdown_to_docx():
    input_md = select_markdown_file()
    if not input_md:
        print("未选择文件，程序已退出。")
        return

    base_name = os.path.splitext(os.path.basename(input_md))[0]
    output_docx = os.path.join(os.path.dirname(input_md), base_name + ".docx")
    markdown_to_word(input_md, output_docx)


if __name__ == "__main__":
    run_markdown_to_docx()
