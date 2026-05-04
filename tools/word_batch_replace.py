import os
import sys

from docx import Document
from PySide6.QtWidgets import (
    QApplication,
    QDialog,
    QDialogButtonBox,
    QFileDialog,
    QFormLayout,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QMessageBox,
    QProgressBar,
    QPushButton,
    QVBoxLayout,
)


def replace_text_in_paragraph(paragraph, old_text, new_text):
    runs = paragraph.runs
    full_text = "".join(run.text for run in runs)

    if old_text in full_text:
        if runs:
            font_name = runs[0].font.name
            font_size = runs[0].font.size
            bold = runs[0].bold
            italic = runs[0].italic
            underline = runs[0].underline

        full_text = full_text.replace(old_text, new_text)
        paragraph.clear()
        new_run = paragraph.add_run(full_text)

        if runs:
            if font_name:
                new_run.font.name = font_name
            if font_size:
                new_run.font.size = font_size
            new_run.bold = bold
            new_run.italic = italic
            new_run.underline = underline


def replace_text_in_headers_footers(doc, old_text, new_text):
    modified = False
    for section in doc.sections:
        header = section.header
        if header.is_linked_to_previous:
            continue
        for paragraph in header.paragraphs:
            before = paragraph.text
            replace_text_in_paragraph(paragraph, old_text, new_text)
            if before != paragraph.text:
                modified = True

        footer = section.footer
        if footer.is_linked_to_previous:
            continue
        for paragraph in footer.paragraphs:
            before = paragraph.text
            replace_text_in_paragraph(paragraph, old_text, new_text)
            if before != paragraph.text:
                modified = True
    return modified


def replace_text_in_docx(docx_file, old_text, new_text):
    doc = Document(docx_file)
    modified = False

    for para in doc.paragraphs:
        before = para.text
        replace_text_in_paragraph(para, old_text, new_text)
        if para.text != before:
            modified = True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    before = para.text
                    replace_text_in_paragraph(para, old_text, new_text)
                    if para.text != before:
                        modified = True

    modified |= replace_text_in_headers_footers(doc, old_text, new_text)

    if modified:
        os.chmod(docx_file, 0o666)
        doc.save(docx_file)


def search_and_replace(root_folder, old_text, new_text):
    count = 0
    for root, dirs, files in os.walk(root_folder):
        for file in files:
            if file.endswith(".docx"):
                docx_file = os.path.join(root, file)
                try:
                    replace_text_in_docx(docx_file, old_text, new_text)
                    count += 1
                except Exception as e:
                    print(f"处理 {docx_file} 失败: {e}")
    return count


def _collect_docx_files(root_folder: str) -> list[str]:
    matches = []
    for root, _, files in os.walk(root_folder):
        for file in files:
            if file.endswith(".docx"):
                matches.append(os.path.join(root, file))
    return matches


class _BatchReplaceDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Word 内容批量替换")
        self.setMinimumWidth(460)
        layout = QVBoxLayout(self)

        form = QFormLayout()

        # 文件夹
        folder_row = QHBoxLayout()
        self._folder_input = QLineEdit()
        self._folder_input.setPlaceholderText("选择包含 Word 文件的文件夹…")
        browse_btn = QPushButton("浏览")
        browse_btn.clicked.connect(self._browse_folder)
        folder_row.addWidget(self._folder_input)
        folder_row.addWidget(browse_btn)
        form.addRow("目标文件夹：", folder_row)

        # 查找 / 替换
        self._old_text = QLineEdit()
        self._new_text = QLineEdit()
        form.addRow("查找文本：", self._old_text)
        form.addRow("替换为：", self._new_text)
        layout.addLayout(form)

        # --- 确定 / 取消 ---
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.button(QDialogButtonBox.Ok).setText("确定")
        buttons.button(QDialogButtonBox.Cancel).setText("取消")
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

        self._progress = QProgressBar()
        self._progress.setRange(0, 1)
        self._progress.setValue(0)
        layout.addWidget(self._progress)
        self._progress.hide()

    def _browse_folder(self):
        d = QFileDialog.getExistingDirectory(self, "选择文件夹")
        if d:
            self._folder_input.setText(d)

    def folder(self) -> str:
        return self._folder_input.text().strip()

    def old_text(self) -> str:
        return self._old_text.text()

    def new_text(self) -> str:
        return self._new_text.text()

    def start_progress(self, maximum: int) -> None:
        self._progress.setRange(0, max(1, maximum))
        self._progress.setValue(0)
        self.show()
        QApplication.processEvents()

    def set_progress(self, value: int) -> None:
        self._progress.setValue(value)
        QApplication.processEvents()


def run_batch_replace(progress_callback=None) -> str:
    app = QApplication.instance() or QApplication(sys.argv)

    dialog = _BatchReplaceDialog()
    if dialog.exec() != QDialog.Accepted:
        return "已取消：Word 内容批量替换。"

    folder = dialog.folder()
    if not folder:
        QMessageBox.warning(None, "提示", "未选择文件夹。")
        return

    old_text = dialog.old_text()
    if old_text is None:
        return

    new_text = dialog.new_text()

    try:
        files = _collect_docx_files(folder)
        if progress_callback:
            progress_callback(0, len(files))
        processed = 0
        for index, file_path in enumerate(files, start=1):
            replace_text_in_docx(file_path, old_text, new_text)
            processed += 1
            if progress_callback:
                progress_callback(index, len(files))
        return f"Word 内容批量替换完成，共处理 {processed} 个文件。"
    except Exception as exc:
        QMessageBox.critical(None, "替换失败", str(exc))
        return f"Word 内容批量替换失败：{exc}"


if __name__ == "__main__":
    run_batch_replace()
